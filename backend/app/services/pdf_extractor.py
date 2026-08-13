from pathlib import Path
import re

import fitz  # PyMuPDF
from docx import Document


class PDFService:
    """
    Service responsible for extracting and cleaning text from PDF and DOCX files.

    Supports:
        - PDF
        - DOCX

    Notes:
        - PDF text extraction uses PyMuPDF.
        - DOCX extraction includes paragraphs and tables.
        - Extracted text is normalized to reduce encoding/unusual-character
          problems before being passed to downstream AI processing.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def extract_text(self, file_path: str | Path) -> str:
        """
        Extract and clean text from a PDF or DOCX file.

        Args:
            file_path: Path to the document.

        Returns:
            Cleaned extracted text.

        Raises:
            FileNotFoundError:
                If the file does not exist.
            ValueError:
                If the file type is unsupported or no readable text was found.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        if not path.is_file():
            raise ValueError(f"Path is not a file: {path}")

        extension = path.suffix.lower()

        if extension == ".pdf":
            text = self._extract_pdf(path)
        elif extension == ".docx":
            text = self._extract_docx(path)
        else:
            raise ValueError(
                f"Unsupported file type '{extension}'. "
                f"Supported types: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        text = self._clean_text(text)

        if not text:
            raise ValueError(
                "No readable text could be extracted from the document. "
                "The file may be empty, image-based, scanned, or corrupted."
            )

        return text

    def _extract_pdf(self, path: Path) -> str:
        """
        Extract text from every page of a PDF.

        Note:
            This works for text-based PDFs. Scanned/image-only PDFs require OCR.
        """
        pages = []

        try:
            with fitz.open(path) as document:
                for page in document:
                    page_text = page.get_text("text")

                    if page_text and page_text.strip():
                        pages.append(page_text)

        except Exception as exc:
            raise ValueError(f"Failed to read PDF '{path}': {exc}") from exc

        return "\n".join(pages)

    def _extract_docx(self, path: Path) -> str:
        """
        Extract text from DOCX paragraphs and tables.
        """
        try:
            document = Document(path)
        except Exception as exc:
            raise ValueError(f"Failed to read DOCX '{path}': {exc}") from exc

        parts = []

        # Normal paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        cells.append(cell_text)

                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Normalize extracted text so unusual encoding artifacts and invisible
        characters are less likely to cause problems downstream.

        Important:
            This does NOT remove normal international characters such as
            Arabic letters, accented characters, etc.
        """

        if not text:
            return ""

        # Normalize line endings.
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Remove Unicode BOM if it appears in extracted text.
        text = text.replace("\ufeff", "")

        # Replace non-breaking spaces and similar spacing characters.
        text = text.replace("\u00a0", " ")

        # Replace common Unicode dash variants with a normal hyphen.
        text = re.sub(r"[\u2010\u2011\u2012\u2013\u2014\u2212]", "-", text)

        # Replace common Unicode quote variants with ASCII equivalents.
        text = re.sub(r"[\u2018\u2019\u201a\u201b]", "'", text)
        text = re.sub(r"[\u201c\u201d\u201e\u201f]", '"', text)

        # Replace ellipsis with three normal dots.
        text = text.replace("\u2026", "...")

        # Remove zero-width/invisible formatting characters.
        text = re.sub(r"[\u200b\u200c\u200d\u2060]", "", text)

        # Remove control characters while preserving:
        # \n = newline
        # \t = tab
        text = "".join(
            char
            for char in text
            if char in "\n\t" or not ord(char) < 32
        )

        # Remove Unicode directional formatting marks.
        text = re.sub(r"[\u202a-\u202e\u2066-\u2069]", "", text)

        # Clean excessive spaces while preserving line structure.
        text = re.sub(r"[ \t]+", " ", text)

        # Remove excessive blank lines.
        text = re.sub(r"\n[ \t]*\n(?:[ \t]*\n)+", "\n\n", text)

        return text.strip()
